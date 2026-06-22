#!/usr/bin/env python3
"""Generate thesis docx file with proper formatting."""
import os
import sys
import json

# Ensure python-docx
try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml
except ImportError:
    os.system('pip install python-docx')
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml

LQ = '\u201c'
RQ = '\u201d'

def sp(doc, text, fn=None, fs=None, bold=False, align=None, indent=True, sb=None, sa=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    if bold: r.bold = True
    if fn:
        r.font.name = 'Times New Roman'
        r._element.rPr.rFonts.set(qn('w:eastAsia'), fn)
        r._element.rPr.rFonts.set(qn('w:ascii'), 'Times New Roman')
        r._element.rPr.rFonts.set(qn('w:hAnsi'), 'Times New Roman')
    else:
        r.font.name = 'Times New Roman'
        r._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
        r._element.rPr.rFonts.set(qn('w:ascii'), 'Times New Roman')
        r._element.rPr.rFonts.set(qn('w:hAnsi'), 'Times New Roman')
    if fs: r.font.size = fs
    if align is not None: p.alignment = align
    p.paragraph_format.first_line_indent = Cm(0.74) if indent else None
    if sb is not None: p.paragraph_format.space_before = sb
    if sa is not None: p.paragraph_format.space_after = sa
    return p

def page_num(doc):
    s = doc.sections[-1]
    for footer, align in [(s.footer, WD_ALIGN_PARAGRAPH.RIGHT), (s.even_page_footer, WD_ALIGN_PARAGRAPH.LEFT)]:
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = align
        r1 = p.add_run()
        r1._r.append(parse_xml('<w:fldChar %s w:fldCharType="begin"/>' % nsdecls('w')))
        r2 = p.add_run()
        r2._r.append(parse_xml('<w:instrText %s xml:space="preserve"> PAGE </w:instrText>' % nsdecls('w')))
        r3 = p.add_run()
        r3._r.append(parse_xml('<w:fldChar %s w:fldCharType="end"/>' % nsdecls('w')))

def setup_styles(doc):
    for name, size in [('Heading 1', Pt(15)), ('Heading 2', Pt(14)), ('Heading 3', Pt(12))]:
        h = doc.styles[name]
        h.font.name = 'Times New Roman'
        h.font.size = size
        h.font.bold = True
        h.font.color.rgb = RGBColor(0,0,0)
        h.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')
        h.element.rPr.rFonts.set(qn('w:ascii'), 'Times New Roman')
        h.element.rPr.rFonts.set(qn('w:hAnsi'), 'Times New Roman')
        h.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        h.paragraph_format.first_line_indent = None
        h.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        h.paragraph_format.space_before = Pt(12) if name == 'Heading 1' else Pt(6)
        h.paragraph_format.space_after = Pt(6) if name == 'Heading 1' else Pt(3)
    n = doc.styles['Normal']
    n.font.name = 'Times New Roman'
    n.font.size = Pt(12)
    n.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    n.element.rPr.rFonts.set(qn('w:ascii'), 'Times New Roman')
    n.element.rPr.rFonts.set(qn('w:hAnsi'), 'Times New Roman')
    n.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    n.paragraph_format.first_line_indent = Cm(0.74)

def mk_cover(doc):
    """封面页：严格按格式规范模板，黑体22pt题名、黑体16pt学生信息、无缩进"""
    from docx.oxml import OxmlElement

    def _mk_run(para, text, font_name='Times New Roman', east_asia='\u9ed1\u4f53', size=Pt(16), bold=False):
        run = para.add_run(text)
        run.font.name = font_name
        run.font.size = size
        run.font.bold = bold
        rPr = run._r.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rPr.insert(0, rFonts)
        rFonts.set(qn('w:ascii'), font_name)
        rFonts.set(qn('w:hAnsi'), font_name)
        rFonts.set(qn('w:eastAsia'), east_asia)

    # P0: 空行 (12pt)
    p = doc.add_paragraph('')
    p.paragraph_format.space_before = None

    # P1: 空行 (center)
    p = doc.add_paragraph('')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # P2: 空行 (16pt Bold, 段前16pt)
    p = doc.add_paragraph('')
    _mk_run(p, '', size=Pt(16), bold=True)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(16)

    # P3: "毕业论文" 36pt Bold 居中，黑体
    p = doc.add_paragraph('')
    _mk_run(p, '\u6bd5\u4e1a\u8bba\u6587', east_asia='\u9ed1\u4f53', size=Pt(36), bold=True)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # P4: 空行 (24pt Bold)
    p = doc.add_paragraph('')
    _mk_run(p, '', size=Pt(24), bold=True)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # P5: "题    名：" 黑体22pt
    p = doc.add_paragraph('')
    _mk_run(p, '\u9898    \u540d\uff1a', font_name='\u9ed1\u4f53', east_asia='\u9ed1\u4f53', size=Pt(22))

    # P6: 标题内容 黑体22pt
    p = doc.add_paragraph('')
    _mk_run(p, '\u9762\u5411\u4f01\u4e1a\u5927\u6570\u636e\u7684\u667a\u80fd\u77e5\u8bc6\u5e93\u95ee\u7b54\u7cfb\u7edf\u7684\u8bbe\u8ba1\u4e0e\u5f00\u53d1', font_name='\u9ed1\u4f53', east_asia='\u9ed1\u4f53', size=Pt(22))

    # P7-P9: 3个空行 (16pt)
    for _ in range(3):
        p = doc.add_paragraph('')
        _mk_run(p, '', size=Pt(16))
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # P10-P14: 学生信息，黑体16pt
    for t in ['\u5b66\u751f\u59d3\u540d\uff1a\u6c5f\u653f\u5bbe',
              '\u5b66   \u9662\uff1a\u4eba\u5de5\u667a\u80fd\u5b66\u9662',
              '\u73ed   \u7ea7\uff1a22511',
              '\u6307\u5bfc\u6559\u5e08\uff1a\u80e1\u529b\u6587',
              '\u5b8c\u6210\u65e5\u671f\uff1a 2026  \u5e74  6  \u6708  20  \u65e5']:
        p = doc.add_paragraph('')
        _mk_run(p, t, size=Pt(16))
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # P15: 空行 (16pt)
    p = doc.add_paragraph('')
    _mk_run(p, '', size=Pt(16))
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    doc.add_page_break()

def mk_abstract(doc):
    abs_title = doc.add_heading('摘  要', level=1)
    abs_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('')
    sp(doc, '\u968f\u7740\u4f01\u4e1a\u4fe1\u606f\u5316\u5efa\u8bbe\u7684\u6301\u7eed\u63a8\u8fdb\uff0c\u5404\u7c7b\u4e1a\u52a1\u7cfb\u7edf\u79ef\u7d2f\u4e86\u89c4\u6a21\u5e9e\u5927\u7684\u6570\u636e\u8d44\u6e90\uff0c\u6db5\u76d6\u6587\u6863\u8d44\u6599\u3001\u7ed3\u6784\u5316\u62a5\u8868\u4ee5\u53ca\u6570\u636e\u5e93\u8bb0\u5f55\u7b49\u591a\u79cd\u5f62\u5f0f\u3002\u7136\u800c\uff0c\u8fd9\u4e9b\u6570\u636e\u5f80\u5f80\u5206\u6563\u5728\u4e0d\u540c\u7684\u7cfb\u7edf\u5e73\u53f0\u4e2d\uff0c\u5f62\u6210\u4e86\u4fe1\u606f\u5b64\u5c9b\uff0c\u5458\u5de5\u5728\u67e5\u627e\u7279\u5b9a\u4e1a\u52a1\u77e5\u8bc6\u65f6\u9700\u8981\u8de8\u8d8a\u591a\u4e2a\u7cfb\u7edf\uff0c\u6548\u7387\u4f4e\u4e0b\u4e14\u5bb9\u6613\u9057\u6f0f\u5173\u952e\u4fe1\u606f\u3002\u4f20\u7edf\u7684\u5173\u952e\u8bcd\u68c0\u7d22\u65b9\u5f0f\u96be\u4ee5\u7406\u89e3\u7528\u6237\u7684\u771f\u5b9e\u67e5\u8be2\u610f\u56fe\uff0c\u9762\u5bf9\u8bed\u4e49\u76f8\u8fd1\u4f46\u8868\u8ff0\u4e0d\u540c\u7684\u95ee\u9898\u65f6\uff0c\u68c0\u7d22\u6548\u679c\u5f80\u5f80\u4e0d\u7406\u60f3\u3002')
    sp(doc, '\u9488\u5bf9\u4e0a\u8ff0\u95ee\u9898\uff0c\u672c\u6587\u8bbe\u8ba1\u5e76\u5b9e\u73b0\u4e86\u4e00\u5957\u9762\u5411\u4f01\u4e1a\u5927\u6570\u636e\u573a\u666f\u7684\u667a\u80fd\u77e5\u8bc6\u5e93\u95ee\u7b54\u7cfb\u7edf\u3002\u8be5\u7cfb\u7edf\u4ee5\u68c0\u7d22\u589e\u5f3a\u751f\u6210\uff08Retrieval-Augmented Generation\uff0cRAG\uff09\u6280\u672f\u4e3a\u6838\u5fc3\u67b6\u6784\uff0c\u6574\u5408\u4e86\u591a\u6e90\u6570\u636e\u63a5\u5165\u3001\u81ea\u52a8\u5316\u6570\u636e\u5904\u7406\u3001\u8bed\u4e49\u5411\u91cf\u68c0\u7d22\u4ee5\u53ca\u5927\u8bed\u8a00\u6a21\u578b\u63a8\u7406\u7b49\u5173\u952e\u6280\u672f\uff0c\u5b9e\u73b0\u4e86\u4ece\u539f\u59cb\u6570\u636e\u5230\u7cbe\u51c6\u95ee\u7b54\u7684\u5b8c\u6574\u5904\u7406\u94fe\u8def\u3002\u5728\u6570\u636e\u5904\u7406\u5c42\u9762\uff0c\u7cfb\u7edf\u652f\u6301\u4e09\u79cd\u5904\u7406\u5f15\u64ce\uff1a\u5feb\u901f\u5165\u5e93\u5f15\u64ce\u7528\u4e8e\u65e5\u5e38\u6587\u6863\u7684\u5373\u65f6\u89e3\u6790\uff0c\u6279\u5904\u7406\u5f15\u64ce\u501f\u52a9 Apache Spark \u5b9e\u73b0\u6d77\u91cf\u6570\u636e\u7684\u5206\u5e03\u5f0f\u6e05\u6d17\u4e0e\u53bb\u91cd\uff0c\u6570\u636e\u5e93\u5f15\u64ce\u901a\u8fc7 Text-to-SQL \u6280\u672f\u5c06\u7528\u6237\u7684\u81ea\u7136\u8bed\u8a00\u67e5\u8be2\u8f6c\u6362\u4e3a\u7ed3\u6784\u5316\u67e5\u8be2\u8bed\u53e5\uff0c\u76f4\u63a5\u4ece\u5173\u7cfb\u578b\u6570\u636e\u5e93\u4e2d\u83b7\u53d6\u7cbe\u786e\u7ed3\u679c\u3002\u5728\u68c0\u7d22\u5c42\u9762\uff0c\u7cfb\u7edf\u91c7\u7528 BGE-large-zh-v1.5 \u6a21\u578b\u5c06\u6587\u672c\u5207\u7247\u6620\u5c04\u4e3a1024\u7ef4\u7a20\u5bc6\u5411\u91cf\uff0c\u5b58\u5165 Milvus \u5411\u91cf\u6570\u636e\u5e93\uff0c\u5e76\u7ed3\u5408\u6df7\u5408\u68c0\u7d22\u7b56\u7565\u540c\u65f6\u5229\u7528\u8bed\u4e49\u76f8\u4f3c\u5ea6\u548c\u5173\u952e\u8bcd\u5339\u914d\u6765\u63d0\u5347\u53ec\u56de\u8d28\u91cf\u3002\u5728\u751f\u6210\u5c42\u9762\uff0c\u7cfb\u7edf\u901a\u8fc7\u591a\u5c42\u67e5\u8be2\u8def\u7531\u673a\u5236\u81ea\u52a8\u8bc6\u522b\u7528\u6237\u95ee\u9898\u7c7b\u578b\uff0c\u5206\u522b\u8c03\u7528 RAG \u68c0\u7d22\u3001Text-to-SQL\u3001\u901a\u7528\u77e5\u8bc6\u6216\u95f2\u804a\u7b49\u4e0d\u540c\u5904\u7406\u8def\u5f84\uff0c\u7531\u5927\u8bed\u8a00\u6a21\u578b\u57fa\u4e8e\u68c0\u7d22\u7ed3\u679c\u751f\u6210\u81ea\u7136\u6d41\u7545\u7684\u56de\u7b54\u3002')
    sp(doc, '\u7cfb\u7edf\u91c7\u7528\u524d\u540e\u7aef\u5206\u79bb\u7684\u6280\u672f\u67b6\u6784\uff0c\u540e\u7aef\u57fa\u4e8e FastAPI \u6846\u67b6\u6784\u5efa RESTful API \u670d\u52a1\uff0c\u524d\u7aef\u91c7\u7528 Vue 3 \u7ed3\u5408 Element Plus \u7ec4\u4ef6\u5e93\u5b9e\u73b0\u7528\u6237\u4ea4\u4e92\u754c\u9762\u3002\u7ecf\u5b9e\u9645\u6d4b\u8bd5\u9a8c\u8bc1\uff0c\u7cfb\u7edf\u80fd\u591f\u6709\u6548\u5904\u7406\u591a\u79cd\u683c\u5f0f\u7684\u4f01\u4e1a\u6587\u6863\uff0c\u67e5\u8be2\u8def\u7531\u7684\u5206\u7c7b\u51c6\u786e\u7387\u8fbe\u5230\u4e86\u9884\u671f\u6c34\u5e73\uff0cText-to-SQL \u529f\u80fd\u53ef\u6b63\u786e\u5904\u7406\u5e38\u89c1\u7684\u6570\u636e\u5206\u6790\u67e5\u8be2\uff0c\u6574\u4f53\u54cd\u5e94\u65f6\u95f4\u63a7\u5236\u5728\u7528\u6237\u53ef\u63a5\u53d7\u7684\u8303\u56f4\u5185\u3002\u672c\u7cfb\u7edf\u4e3a\u4f01\u4e1a\u77e5\u8bc6\u7ba1\u7406\u63d0\u4f9b\u4e86\u4e00\u79cd\u53ef\u884c\u7684\u6280\u672f\u65b9\u6848\uff0c\u5177\u6709\u4e00\u5b9a\u7684\u5b9e\u9645\u5e94\u7528\u4ef7\u503c\u3002')
    doc.add_paragraph('')
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74)
    r1 = p.add_run('\u5173\u952e\u8bcd\uff1a')
    r1.bold = True
    r1.font.name = 'Times New Roman'
    r1.font.size = Pt(12)
    r1._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    r1._element.rPr.rFonts.set(qn('w:ascii'), 'Times New Roman')
    r1._element.rPr.rFonts.set(qn('w:hAnsi'), 'Times New Roman')
    r2 = p.add_run('\u77e5\u8bc6\u5e93\u95ee\u7b54\u7cfb\u7edf\uff1b\u68c0\u7d22\u589e\u5f3a\u751f\u6210\uff1b\u5927\u8bed\u8a00\u6a21\u578b\uff1bText-to-SQL\uff1b\u5411\u91cf\u68c0\u7d22')
    r2.font.name = 'Times New Roman'
    r2.font.size = Pt(12)
    r2._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    r2._element.rPr.rFonts.set(qn('w:ascii'), 'Times New Roman')
    r2._element.rPr.rFonts.set(qn('w:hAnsi'), 'Times New Roman')
    page_num(doc)
    doc.add_page_break()

def mk_toc(doc):
    toc_title = doc.add_heading('\u76ee  \u5f55', level=1)
    toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para = doc.add_paragraph()
    para.paragraph_format.first_line_indent = None
    r = para.add_run()
    r._r.append(parse_xml('<w:fldChar %s w:fldCharType="begin"/>' % nsdecls('w')))
    r2 = para.add_run()
    r2._r.append(parse_xml('<w:instrText %s xml:space="preserve"> TOC \\o "1-3" \\h \\z \\u </w:instrText>' % nsdecls('w')))
    r3 = para.add_run()
    r3._r.append(parse_xml('<w:fldChar %s w:fldCharType="separate"/>' % nsdecls('w')))
    r4 = para.add_run('(\u8bf7\u5728 Word \u4e2d\u53f3\u952e\u6b64\u5904\uff0c\u9009\u62e9' + LQ + '\u66f4\u65b0\u57df' + RQ + '\u4ee5\u751f\u6210\u76ee\u5f55)')
    r4.font.color.rgb = RGBColor(128,128,128)
    r4.font.size = Pt(10)
    r5 = para.add_run()
    r5._r.append(parse_xml('<w:fldChar %s w:fldCharType="end"/>' % nsdecls('w')))
    doc.add_page_break()

def main():
    # Load content
    content_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'thesis_content.json')
    with open(content_path, 'r', encoding='utf-8') as f:
        C = json.load(f)

    doc = Document()
    s = doc.sections[0]
    s.page_width = Cm(21)
    s.page_height = Cm(29.7)
    s.top_margin = Cm(2.54)
    s.bottom_margin = Cm(2.54)
    s.left_margin = Cm(3.17)
    s.right_margin = Cm(3.17)
    s._sectPr.append(parse_xml('<w:evenAndOddHeaders %s w:val="true"/>' % nsdecls('w')))

    setup_styles(doc)
    mk_cover(doc)
    mk_abstract(doc)
    mk_toc(doc)

    # Chapters
    for i, ch in enumerate(C['chapters'], 1):
        doc.add_heading(f'{i}. {ch["title"].split(" ", 1)[-1] if " " in ch["title"] else ch["title"]}', level=1)
        for sec in ch.get('sections', []):
            doc.add_heading(sec['title'], level=2)
            for sub in sec.get('subsections', []):
                if sub.get('level') == 3:
                    doc.add_heading(sub['title'], level=3)
                for para in sub.get('paragraphs', []):
                    sp(doc, para)
            for para in sec.get('paragraphs', []):
                sp(doc, para)

    # Conclusion
    num = len(C['chapters']) + 1
    doc.add_heading(f'{num}. {C["conclusion"]["title"]}', level=1)
    for para in C['conclusion']['paragraphs']:
        sp(doc, para)

    # Acknowledgement
    doc.add_page_break()
    ack_title = doc.add_heading(C['acknowledgement']['title'], level=1)
    ack_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('')
    for para in C['acknowledgement']['paragraphs']:
        sp(doc, para)

    # References
    doc.add_page_break()
    ref_title = doc.add_heading(C['references']['title'], level=1)
    ref_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('')
    for ref in C['references']['items']:
        p = doc.add_paragraph()
        r = p.add_run(ref)
        # 规范要求：宋体五号(10.5pt)，全部字体统一宋体（含英文）
        r.font.name = '宋体'
        r.font.size = Pt(10.5)
        r._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        r._element.rPr.rFonts.set(qn('w:ascii'), '宋体')
        r._element.rPr.rFonts.set(qn('w:hAnsi'), '宋体')
        # 首行缩进2字符（非悬挂缩进）
        p.paragraph_format.first_line_indent = Cm(0.74)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

    # Save
    out = os.path.join(os.path.dirname(os.path.abspath(__file__)), C['output_filename'])
    doc.save(out)
    print('Generated: ' + out)
    print('Remember: Ctrl+A then F9 to update TOC in Word')

if __name__ == '__main__':
    main()
